"""Verify DOCX has embedded images and working bookmarks."""
from docx import Document
import os

doc = Document('PROJECT_DOCUMENTATION.docx')

# Count images
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1

# Count bookmarks
bookmark_count = 0
from docx.oxml.ns import qn
for p in doc.paragraphs:
    for elem in p._element.iter():
        if elem.tag == qn('w:bookmarkStart'):
            bookmark_count += 1

# Count hyperlinks
hyperlink_count = 0
for p in doc.paragraphs:
    for elem in p._element.iter():
        if elem.tag == qn('w:hyperlink'):
            anchor = elem.get(qn('w:anchor'))
            if anchor:
                hyperlink_count += 1

# Count headings
heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))

# Count paragraphs and tables
print(f"DOCX Verification Report")
print(f"========================")
print(f"File size:      {os.path.getsize('PROJECT_DOCUMENTATION.docx'):,} bytes")
print(f"Paragraphs:     {len(doc.paragraphs)}")
print(f"Tables:         {len(doc.tables)}")
print(f"Headings:       {heading_count}")
print(f"Images:         {image_count}")
print(f"Bookmarks:      {bookmark_count}")
print(f"TOC Hyperlinks: {hyperlink_count}")

# Check first few hyperlinks
print(f"\nFirst TOC hyperlinks found:")
count = 0
for p in doc.paragraphs:
    for elem in p._element.iter():
        if elem.tag == qn('w:hyperlink'):
            anchor = elem.get(qn('w:anchor'))
            if anchor and count < 5:
                print(f"  → #{anchor}")
                count += 1
