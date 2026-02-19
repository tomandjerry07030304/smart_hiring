"""
Generate PROJECT_DOCUMENTATION.docx and PROJECT_DOCUMENTATION.pdf
from PROJECT_DOCUMENTATION.md with:
  - Times New Roman font throughout
  - Embedded PNG diagrams from diagrams/ folder (replacing Mermaid blocks)
  - Working Table of Contents with clickable internal hyperlinks
  - Professional tables with dark blue headers and alternating rows
  - Code blocks with Consolas font and gray background
  - Blockquotes with left border (for Q&A section)
"""
import re
import os
import glob
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

FONT_NAME = 'Times New Roman'
INPUT_MD = 'PROJECT_DOCUMENTATION.md'
OUTPUT_DOCX = 'PROJECT_DOCUMENTATION.docx'
OUTPUT_PDF = 'PROJECT_DOCUMENTATION.pdf'
DIAGRAMS_DIR = 'diagrams'

# ──────────────────────────────────────────────────
# Map mermaid block index (1-11) to PNG file
# ──────────────────────────────────────────────────
def get_diagram_map():
    """Build a mapping from diagram index to PNG path."""
    pngs = sorted(glob.glob(os.path.join(DIAGRAMS_DIR, '*.png')))
    mapping = {}
    for png in pngs:
        basename = os.path.basename(png)
        # Extract index from filename like 01_name.png
        match = re.match(r'^(\d+)_', basename)
        if match:
            mapping[int(match.group(1))] = png
    return mapping

# ──────────────────────────────────────────────────
# Style Setup
# ──────────────────────────────────────────────────
def setup_styles(doc):
    """Set up document styles with Times New Roman."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    for i, size in enumerate([24, 20, 16, 14], 1):
        try:
            hs = doc.styles[f'Heading {i}']
        except KeyError:
            hs = doc.styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
        hs.font.name = FONT_NAME
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 51, 102)
        hs.paragraph_format.space_before = Pt(18 if i <= 2 else 12)
        hs.paragraph_format.space_after = Pt(8)
        hs.paragraph_format.keep_with_next = True

    # Code style
    try:
        code_style = doc.styles['Code']
    except KeyError:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor(30, 30, 30)
    code_style.paragraph_format.space_before = Pt(1)
    code_style.paragraph_format.space_after = Pt(1)
    code_style.paragraph_format.line_spacing = 1.0

# ──────────────────────────────────────────────────
# Bookmark / Hyperlink Helpers
# ──────────────────────────────────────────────────
def add_bookmark(paragraph, bookmark_name):
    """Add an invisible bookmark anchor to a paragraph."""
    tag = paragraph._element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) & 0x7FFFFFFF))
    bookmark_start.set(qn('w:name'), bookmark_name)
    tag.insert(0, bookmark_start)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) & 0x7FFFFFFF))
    tag.append(bookmark_end)

def add_internal_hyperlink(paragraph, bookmark_name, display_text):
    """Add a clickable internal hyperlink to an existing bookmark."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), '003366')
    rPr.append(color_elem)

    u_elem = OxmlElement('w:u')
    u_elem.set(qn('w:val'), 'single')
    rPr.append(u_elem)

    font_elem = OxmlElement('w:rFonts')
    font_elem.set(qn('w:ascii'), FONT_NAME)
    font_elem.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(font_elem)

    sz_elem = OxmlElement('w:sz')
    sz_elem.set(qn('w:val'), '22')  # 11pt
    rPr.append(sz_elem)

    run_elem.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.set(qn('xml:space'), 'preserve')
    text_elem.text = display_text
    run_elem.append(text_elem)

    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)

# ──────────────────────────────────────────────────
# Code Block
# ──────────────────────────────────────────────────
def add_code_block(doc, code_text, lang=''):
    """Add a formatted code block with optional language label."""
    if lang:
        p = doc.add_paragraph()
        run = p.add_run(f'  [{lang.upper()}]')
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True

    lines = code_text.split('\n')
    for line in lines:
        p = doc.add_paragraph(style='Code')
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        p._element.get_or_add_pPr().append(shading)

# ──────────────────────────────────────────────────
# Table
# ──────────────────────────────────────────────────
def add_table_from_rows(doc, header_row, data_rows):
    """Add a formatted table with dark blue header and alternating rows."""
    if not header_row:
        return
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            text = cell_text.strip()
            # Handle inline bold
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for pi, part in enumerate(parts):
                if pi % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = FONT_NAME
                run.font.size = Pt(9)

            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # spacing

def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx."""
    header_line = lines[start_idx].strip()
    header = [c.strip() for c in header_line.split('|')[1:-1]]
    end_idx = start_idx + 2  # skip separator
    rows = []
    while end_idx < len(lines):
        line = lines[end_idx].strip()
        if not line.startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        end_idx += 1
    return header, rows, end_idx

# ──────────────────────────────────────────────────
# Inline-Formatted Paragraph
# ──────────────────────────────────────────────────
def add_formatted_paragraph(doc, text, style='Normal', bold=False):
    """Add paragraph with bold/code inline formatting."""
    p = doc.add_paragraph(style=style)
    if bold:
        run = p.add_run(text)
        run.bold = True
        run.font.name = FONT_NAME
        return p

    tokens = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = p.add_run(token[2:-2])
            run.bold = True
            run.font.name = FONT_NAME
        elif token.startswith('`') and token.endswith('`'):
            run = p.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(180, 0, 0)
        else:
            run = p.add_run(token)
            run.font.name = FONT_NAME
    return p

# ──────────────────────────────────────────────────
# Heading → Bookmark-name mapping
# ──────────────────────────────────────────────────
def heading_to_bookmark(text):
    """Convert heading text to a bookmark name (matching MD anchor rules)."""
    # Remove special chars, lowercase, replace spaces with hyphens
    bm = text.lower().strip()
    bm = re.sub(r'[^\w\s-]', '', bm)             # remove non-word chars except hyphens
    bm = re.sub(r'\s+', '-', bm)                  # spaces → hyphens
    bm = re.sub(r'-+', '-', bm).strip('-')        # collapse hyphens
    return bm

# ──────────────────────────────────────────────────
# MAIN CONVERSION
# ──────────────────────────────────────────────────
def convert_md_to_docx(md_path, docx_path):
    """Convert markdown file to DOCX with embedded diagrams and working TOC links."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    diagram_map = get_diagram_map()
    print(f"   Diagram images found: {len(diagram_map)}")

    lines = content.split('\n')
    doc = Document()
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    i = 0
    in_code_block = False
    code_buffer = []
    code_lang = ''
    mermaid_counter = 0  # track which mermaid block we're on

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code Blocks ──
        if stripped.startswith('```') and not in_code_block:
            in_code_block = True
            code_lang = stripped[3:].strip()
            code_buffer = []
            i += 1
            continue
        elif stripped.startswith('```') and in_code_block:
            in_code_block = False
            if code_lang.lower() == 'mermaid':
                mermaid_counter += 1
                # Try to insert the rendered PNG
                png_path = diagram_map.get(mermaid_counter)
                if png_path and os.path.exists(png_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(png_path, width=Inches(5.5))
                    # Add caption
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(f'Figure {mermaid_counter}')
                    cap_run.font.name = FONT_NAME
                    cap_run.font.size = Pt(9)
                    cap_run.italic = True
                    cap_run.font.color.rgb = RGBColor(80, 80, 80)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(f'[Diagram {mermaid_counter} — see Markdown source for Mermaid rendering]')
                    run.font.name = FONT_NAME
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_code_block(doc, '\n'.join(code_buffer), code_lang)
            code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ── Horizontal Rules ──
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="003366"/>'
                f'</w:pBdr>'
            )
            p._element.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # ── Tables ──
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            header, rows, end_idx = parse_table(lines, i)
            add_table_from_rows(doc, header, rows)
            i = end_idx
            continue

        # ── Headings (with bookmarks) ──
        heading_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            p = doc.add_heading(text, level=min(level, 4))
            for run in p.runs:
                run.font.name = FONT_NAME
            # Add bookmark anchor for TOC navigation
            bm_name = heading_to_bookmark(text)
            add_bookmark(p, bm_name)
            i += 1
            continue

        # ── TOC lines: convert [text](#anchor) to clickable hyperlinks ──
        toc_match = re.match(r'^(\d+)\.\s+\[(.+?)\]\(#(.+?)\)\s*$', stripped)
        if toc_match:
            number = toc_match.group(1)
            display_text = toc_match.group(2)
            anchor = toc_match.group(3)
            p = doc.add_paragraph()
            # Add number
            run_num = p.add_run(f'{number}. ')
            run_num.font.name = FONT_NAME
            run_num.font.size = Pt(11)
            # Add clickable hyperlink
            add_internal_hyperlink(p, anchor, display_text)
            i += 1
            continue

        # ── Blockquotes (Q&A answers) ──
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="12" w:space="8" w:color="003366"/>'
                f'</w:pBdr>'
            )
            p._element.get_or_add_pPr().append(pBdr)

            tokens = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
            for token in tokens:
                if token.startswith('**') and token.endswith('**'):
                    run = p.add_run(token[2:-2])
                    run.bold = True
                    run.font.name = FONT_NAME
                elif token.startswith('`') and token.endswith('`'):
                    run = p.add_run(token[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    run = p.add_run(token)
                    run.font.name = FONT_NAME
                for r in p.runs[-1:]:
                    r.font.size = Pt(10)
            i += 1
            continue

        # ── Ordered list ──
        if re.match(r'^\d+\.', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            add_formatted_paragraph(doc, text, style='List Number')
            i += 1
            continue

        # ── Unordered list ──
        if stripped.startswith('- '):
            text = stripped[2:]
            add_formatted_paragraph(doc, text, style='List Bullet')
            i += 1
            continue

        # ── Empty lines ──
        if not stripped:
            i += 1
            continue

        # ── Regular paragraphs ──
        add_formatted_paragraph(doc, stripped)
        i += 1

    doc.save(docx_path)
    print(f"   ✅ Generated {docx_path} ({os.path.getsize(docx_path):,} bytes)")
    return docx_path

# ──────────────────────────────────────────────────
# PDF Conversion
# ──────────────────────────────────────────────────
def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using available methods."""
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"   ✅ Generated {pdf_path} ({os.path.getsize(pdf_path):,} bytes)")
        return True
    except ImportError:
        pass
    try:
        import subprocess
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', '.', docx_path],
            capture_output=True, text=True, timeout=60
        )
        if os.path.exists(pdf_path):
            print(f"   ✅ Generated {pdf_path} ({os.path.getsize(pdf_path):,} bytes)")
            return True
    except (FileNotFoundError, Exception):
        pass
    try:
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        wdoc = word.Documents.Open(os.path.abspath(docx_path))
        wdoc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        wdoc.Close()
        word.Quit()
        print(f"   ✅ Generated {pdf_path} ({os.path.getsize(pdf_path):,} bytes)")
        return True
    except Exception:
        pass
    print(f"   ⚠️  PDF conversion failed. Open {docx_path} in Word → Save As PDF")
    return False

# ──────────────────────────────────────────────────
if __name__ == '__main__':
    print("=" * 60)
    print("Smart Hiring System — Documentation Generator v2")
    print("  • Embedded Mermaid diagram images")
    print("  • Working Table of Contents hyperlinks")
    print("  • Times New Roman + professional formatting")
    print("=" * 60)

    print(f"\n📄 Converting {INPUT_MD} → {OUTPUT_DOCX}...")
    convert_md_to_docx(INPUT_MD, OUTPUT_DOCX)

    print(f"\n📄 Converting {OUTPUT_DOCX} → {OUTPUT_PDF}...")
    convert_docx_to_pdf(OUTPUT_DOCX, OUTPUT_PDF)

    print("\n" + "=" * 60)
    print("✅ Documentation generation complete!")
    print("=" * 60)
