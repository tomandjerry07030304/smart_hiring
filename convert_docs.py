"""
Convert Markdown Documentation to DOCX and PDF formats
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2

def convert_md_to_docx(md_file, output_dir="Documentation_Package"):
    """Convert markdown file to DOCX with proper formatting"""
    print(f"Converting {md_file} to DOCX...")
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse markdown and add to document
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Skip empty lines
        if not line:
            doc.add_paragraph()
            continue
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Handle code blocks
        elif line.startswith('```'):
            continue  # Skip code fence markers
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle numbered lists
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            p = doc.add_paragraph(line[3:], style='List Number')
        
        # Handle bold text markers
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    p.add_run(part).bold = True
        
        # Regular paragraph
        else:
            p = doc.add_paragraph(line)
    
    # Save document
    output_file = os.path.join(output_dir, os.path.basename(md_file).replace('.md', '.docx'))
    doc.save(output_file)
    print(f"✅ Created: {output_file}")
    return output_file

def main():
    """Convert all 7 documentation files"""
    docs = [
        "COMPLETE_PROJECT_ANALYSIS_AND_DOCUMENTATION.md",
        "QUICK_REFERENCE_GUIDE.md",
        "PROJECT_ANALYSIS_SUMMARY_FOR_FACULTY.md",
        "ARCHITECTURE_DIAGRAMS.md",
        "DOCUMENTATION_INDEX.md",
        "DETAILED_CODE_ANALYSIS_APPENDIX.md",
        "TESTING_AND_DEPLOYMENT_GUIDE.md"
    ]
    
    # Create output directory
    output_dir = "Documentation_Package_DOCX"
    os.makedirs(output_dir, exist_ok=True)
    
    print("=" * 60)
    print("Converting Documentation to DOCX Format")
    print("=" * 60)
    
    converted_files = []
    for doc in docs:
        if os.path.exists(doc):
            try:
                output_file = convert_md_to_docx(doc, output_dir)
                converted_files.append(output_file)
            except Exception as e:
                print(f"❌ Error converting {doc}: {str(e)}")
        else:
            print(f"⚠️ File not found: {doc}")
    
    print("\n" + "=" * 60)
    print(f"✅ Conversion Complete! {len(converted_files)}/{len(docs)} files converted")
    print(f"📁 Output directory: {os.path.abspath(output_dir)}")
    print("=" * 60)
    
    # Create ZIP archive
    import zipfile
    zip_name = "Smart_Hiring_Documentation_DOCX.zip"
    with zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in converted_files:
            zipf.write(file, os.path.basename(file))
    
    print(f"\n📦 Created ZIP archive: {zip_name}")
    print(f"📊 Size: {os.path.getsize(zip_name) / 1024:.1f} KB")

if __name__ == "__main__":
    main()
