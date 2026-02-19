"""
Advanced Resume Parser with NLP - Entrepreneur Edition
=========================================================
Enterprise-grade resume parsing with layout-aware extraction

Features:
- Layout-aware PDF extraction (pdfplumber + PyPDF2 fallback)
- Intelligent section detection (Skills, Experience, Education, Projects)
- Context-aware skill extraction (only from relevant sections)
- 350+ skill taxonomy with 15 categories
- Date-range experience calculation (dateparser)
- NLP-based information extraction (spaCy)
- Project & language proficiency extraction  
- Enhanced PII anonymization (DOB, religion, nationality, marital)
- Per-field confidence scoring (0.0-1.0)
- Certification detection
- Backward-compatible with all v3.0 callers

Author: Smart Hiring System Team
Version: 4.0 - Entrepreneur Edition
Date: February 2026
"""

import re
import io
import logging
from typing import Dict, List, Optional, Any, Tuple, Set
from datetime import datetime

from config.skill_ontology import (
    SKILL_CATEGORIES as _ONTOLOGY_CATEGORIES,
    SKILL_DATABASE as _ONTOLOGY_DATABASE,
)

logger = logging.getLogger(__name__)

# =============================================================================
# Dependency Checking with Clear Error Logging
# =============================================================================

# PyPDF2 for PDF extraction
PDF_AVAILABLE = False
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
    logger.info("✅ PyPDF2 loaded successfully - PDF parsing enabled")
except ImportError:
    logger.error(
        "❌ CRITICAL: PyPDF2 not installed. PDF parsing is DISABLED. "
        "Install with: pip install PyPDF2"
    )

# pdfplumber for layout-aware PDF extraction (Entrepreneur Upgrade)
PDFPLUMBER_AVAILABLE = False
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    logger.info("✅ pdfplumber loaded successfully - Layout-aware PDF parsing enabled")
except ImportError:
    logger.warning(
        "⚠️ pdfplumber not installed. Falling back to PyPDF2 for PDF parsing. "
        "Install with: pip install pdfplumber"
    )

# python-docx for DOCX extraction
DOCX_AVAILABLE = False
try:
    from docx import Document
    DOCX_AVAILABLE = True
    logger.info("✅ python-docx loaded successfully - DOCX parsing enabled")
except ImportError:
    logger.error(
        "❌ CRITICAL: python-docx not installed. DOCX parsing is DISABLED. "
        "Install with: pip install python-docx"
    )

# dateparser for intelligent date-range experience calculation
DATEPARSER_AVAILABLE = False
try:
    import dateparser
    DATEPARSER_AVAILABLE = True
    logger.info("✅ dateparser loaded successfully - Smart experience calculation enabled")
except ImportError:
    logger.warning(
        "⚠️ dateparser not installed. Falling back to regex-based experience parsing. "
        "Install with: pip install dateparser"
    )

# spaCy for NLP-based extraction
SPACY_AVAILABLE = False
_spacy_nlp = None
try:
    import spacy
    _spacy_nlp = spacy.load('en_core_web_sm')
    SPACY_AVAILABLE = True
    logger.info("✅ spaCy model 'en_core_web_sm' loaded successfully - NLP extraction enabled")
except ImportError:
    logger.error(
        "❌ CRITICAL: spaCy not installed. NLP-based extraction is DISABLED. "
        "Install with: pip install spacy"
    )
except OSError:
    logger.error(
        "❌ CRITICAL: spaCy model 'en_core_web_sm' not found. NLP-based extraction is DISABLED. "
        "Download with: python -m spacy download en_core_web_sm"
    )

# OCR for scanned PDFs (Gap 9 fix)
OCR_AVAILABLE = False
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    # Set Tesseract path for Windows
    import platform
    if platform.system() == "Windows":
        _tesseract_path = r"C:\Users\venkat anand\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
        import os
        if os.path.exists(_tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = _tesseract_path
    OCR_AVAILABLE = True
    logger.info("✅ pytesseract + pdf2image loaded - OCR for scanned PDFs enabled")
except ImportError:
    logger.warning(
        "⚠️ pytesseract or pdf2image not installed. OCR for scanned PDFs is DISABLED. "
        "Install with: pip install pytesseract pdf2image  (also requires Tesseract binary)"
    )


def _get_nlp():
    """Get spaCy NLP model if available"""
    return _spacy_nlp if SPACY_AVAILABLE else None


# =============================================================================
# Unified Phone Regex - Single source of truth for extraction + anonymization
# =============================================================================

PHONE_PATTERNS = [
    r'\+?\d[\d\-\s()]{6,}\d',                    # International: +1-555-123-4567
    r'(?:\+91|0)?\s?[6-9]\d{9}',                  # Indian mobile: +91 9876543210
    r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US: (555) 123-4567
]


def _merge_date_intervals(intervals):
    """
    Merge overlapping date intervals to prevent double-counting.
    
    Input: [(start_date, end_date), ...]
    Output: [(merged_start, merged_end), ...]
    
    Example:
        [(2020-01, 2022-01), (2021-01, 2023-01)] -> [(2020-01, 2023-01)]
        Result: 3 years instead of 4 (no double-counting)
    """
    if not intervals:
        return []
    
    # Sort by start date
    sorted_intervals = sorted(intervals, key=lambda x: x[0])
    merged = [sorted_intervals[0]]
    
    for current_start, current_end in sorted_intervals[1:]:
        last_start, last_end = merged[-1]
        
        if current_start <= last_end:
            # Overlap detected — extend the interval
            merged[-1] = (last_start, max(last_end, current_end))
        else:
            # No overlap — new interval
            merged.append((current_start, current_end))
    
    return merged


# =============================================================================
# Section Detection - Intelligent Resume Structure Analysis
# =============================================================================

SECTION_HEADERS = {
    'summary': [
        r'(?:professional\s+)?summary', r'objective', r'profile', r'about\s+me',
        r'career\s+(?:summary|objective|profile)', r'personal\s+statement'
    ],
    'experience': [
        r'(?:work|professional|employment)\s+(?:experience|history)',
        r'experience', r'work\s+history', r'career\s+history', r'positions?\s+held'
    ],
    'education': [
        r'education(?:\s+(?:and|&)\s+training)?', r'academic\s+(?:background|qualifications)',
        r'qualifications', r'academic\s+details'
    ],
    'skills': [
        r'(?:technical|core|key|professional)?\s*skills',
        r'technical\s+(?:expertise|proficiency|competencies)',
        r'competencies', r'technologies', r'tools?\s+(?:and|&)\s+technologies'
    ],
    'projects': [
        r'(?:key|major|notable|academic|personal)?\s*projects',
        r'portfolio', r'project\s+experience'
    ],
    'certifications': [
        r'certifications?(?:\s+(?:and|&)\s+licenses?)?',
        r'licenses?\s+(?:and|&)\s+certifications?', r'credentials'
    ],
    'languages': [
        r'languages?(?:\s+(?:known|proficiency))?', r'linguistic\s+skills'
    ],
    'achievements': [
        r'(?:awards?|achievements?|honors?|accomplishments)',
        r'awards?\s+(?:and|&)\s+(?:achievements?|honors?)'
    ],
    'interests': [
        r'(?:hobbies|interests|extracurricular)',
        r'hobbies?\s+(?:and|&)\s+interests?'
    ],
    'declaration': [
        r'declaration', r'disclaimer'
    ],
    'references': [
        r'references?', r'referees?'
    ]
}

# Sections where skills are VALID (context-aware extraction)
SKILL_VALID_SECTIONS = {'skills', 'experience', 'projects', 'summary', 'certifications', None}
# None = unknown section (top of resume before first header) - we allow it

# Sections where skills should be IGNORED
SKILL_IGNORE_SECTIONS = {'interests', 'declaration', 'references'}


def detect_sections(text: str) -> Dict[str, str]:
    """
    Detect and split resume text into labeled sections.
    
    Returns:
        Dict mapping section_name -> section_text content
        Key 'header' contains text before the first detected section.
    """
    lines = text.split('\n')
    sections = {}
    current_section = 'header'
    current_lines = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_lines.append(line)
            continue
        
        # Check if this line is a section header
        detected = None
        for section_name, patterns in SECTION_HEADERS.items():
            for pattern in patterns:
                # Match if the line IS the header (short line, possibly with colon)
                full_pattern = r'^\s*' + pattern + r'\s*:?\s*$'
                if re.match(full_pattern, stripped, re.IGNORECASE):
                    detected = section_name
                    break
            if detected:
                break
        
        if detected:
            # Save previous section
            if current_lines:
                sections[current_section] = '\n'.join(current_lines)
            current_section = detected
            current_lines = []
        else:
            current_lines.append(line)
    
    # Save last section
    if current_lines:
        sections[current_section] = '\n'.join(current_lines)
    
    return sections


# =============================================================================
# Skill Database - Comprehensive 350+ skills with 15 categories
# =============================================================================

# Gap 10: Skills loaded from unified skill_ontology.json (single source of truth)
SKILL_CATEGORIES = _ONTOLOGY_CATEGORIES

# Flat set of all skills for quick lookup
SKILL_DATABASE: Set[str] = _ONTOLOGY_DATABASE

# Education degrees
DEGREES = [
    'phd', 'ph.d', 'doctorate', 'masters', 'master', 'mba', 'ms', 'm.s', 
    'bachelors', 'bachelor', 'bs', 'b.s', 'ba', 'b.a', 'associate', 'diploma',
    'b.tech', 'btech', 'm.tech', 'mtech', 'be', 'b.e', 'me', 'm.e'
]

# Certification patterns
CERTIFICATIONS = [
    'aws certified', 'azure certified', 'google cloud certified', 'cisco',
    'comptia', 'pmp', 'scrum master', 'certified', 'professional', 'specialist',
    'expert', 'associate', 'csm', 'cka', 'ckad'
]


# =============================================================================
# Text Extraction Functions
# =============================================================================

def _extract_pdf_pdfplumber(pdf_file) -> Optional[str]:
    """Extract text using pdfplumber (layout-aware, handles columns)"""
    if not PDFPLUMBER_AVAILABLE:
        return None
    try:
        with pdfplumber.open(pdf_file) as pdf:
            text_parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            full_text = "\n".join(text_parts)
            if full_text.strip() and len(full_text.strip()) >= 50:
                logger.info(f"📄 pdfplumber extracted {len(full_text)} chars")
                return full_text
            return None  # Too short, try fallback
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}, trying fallback")
        return None


def _extract_pdf_pypdf2(pdf_file) -> Optional[str]:
    """Extract text using PyPDF2 (fallback)"""
    if not PDF_AVAILABLE:
        return None
    try:
        reader = PdfReader(pdf_file)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        full_text = "\n".join(text_parts)
        if full_text.strip():
            logger.info(f"📄 PyPDF2 fallback extracted {len(full_text)} chars")
            return full_text
        return None
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
        return None


def _extract_pdf_ocr(raw_bytes: bytes) -> Optional[str]:
    """
    Gap 9: OCR fallback for scanned/image-based PDFs.

    Converts each page to an image, then runs Tesseract OCR.
    Only invoked when pdfplumber and PyPDF2 yield no text.
    """
    if not OCR_AVAILABLE:
        return None
    try:
        images = convert_from_bytes(raw_bytes, dpi=300)
        text_parts = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img)
            if page_text and page_text.strip():
                text_parts.append(page_text)
        full_text = "\n".join(text_parts)
        if full_text.strip() and len(full_text.strip()) >= 30:
            logger.info(f"📄 OCR extracted {len(full_text)} chars from scanned PDF")
            return full_text
        return None
    except Exception as e:
        logger.warning(f"OCR extraction failed: {e}")
        return None


def extract_text_from_pdf(file_data) -> str:
    """
    Extract text from PDF file using multi-engine strategy:
    1. pdfplumber (layout-aware, handles 2-column resumes)
    2. PyPDF2 (fallback for simpler PDFs)
    3. OCR via pytesseract (Gap 9: scanned/image-based PDFs)
    
    Args:
        file_data: Binary PDF file data or file-like object
    
    Returns:
        str: Extracted text content
    """
    if not PDF_AVAILABLE and not PDFPLUMBER_AVAILABLE and not OCR_AVAILABLE:
        logger.error("No PDF parser available")
        return "PDF parsing not available. Install pdfplumber, PyPDF2, or pytesseract."
    
    # Handle both bytes and file-like objects
    if isinstance(file_data, bytes):
        raw_bytes = file_data
    else:
        raw_bytes = file_data.read()
    
    # Strategy 1: pdfplumber (primary - layout-aware)
    result = _extract_pdf_pdfplumber(io.BytesIO(raw_bytes))
    if result:
        return result
    
    # Strategy 2: PyPDF2 (fallback)
    result = _extract_pdf_pypdf2(io.BytesIO(raw_bytes))
    if result:
        return result
    
    # Strategy 3: OCR (Gap 9 — scanned/image-based PDFs)
    result = _extract_pdf_ocr(raw_bytes)
    if result:
        return result
    
    return "No text could be extracted from PDF. The file may be scanned/image-based and OCR is not available."


def extract_text_from_docx(file_data) -> str:
    """
    Extract text from DOCX file (includes paragraphs AND tables)
    
    Args:
        file_data: Binary DOCX file data or file-like object
    
    Returns:
        str: Extracted text content
    """
    if not DOCX_AVAILABLE:
        logger.error("DOCX parsing attempted but python-docx is not installed")
        return "DOCX parsing not available. Please install python-docx."
    
    try:
        # Handle both bytes and file-like objects
        if isinstance(file_data, bytes):
            docx_file = io.BytesIO(file_data)
        else:
            docx_file = file_data
        
        doc = Document(docx_file)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables (important for resumes!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        return full_text if full_text.strip() else "No text could be extracted from DOCX"
    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return f"Error extracting DOCX: {str(e)}"


def extract_text_from_file(file_data, filename: str) -> str:
    """
    Extract text from uploaded file based on extension
    
    Args:
        file_data: Binary file data
        filename: Original filename for format detection
    
    Returns:
        str: Extracted text content
    """
    name = filename.lower()
    
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_data)
    elif name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx(file_data)
    else:
        # Treat as plain text
        try:
            return file_data.decode('utf-8', errors='ignore')
        except Exception:
            return str(file_data)


# =============================================================================
# Anonymization Functions - Remove PII for bias-free hiring
# =============================================================================

def anonymize_text(text: str) -> str:
    """
    Remove PII from text for bias-free resume screening
    
    Removes: emails, phone numbers, URLs, gender indicators, names,
    DOB, age, religion, nationality, marital status
    
    Args:
        text: Original resume text
    
    Returns:
        str: Anonymized text with PII replaced by placeholders
    """
    if not isinstance(text, str) or not text:
        return ""
    
    # Remove emails
    text = re.sub(r'\S+@\S+', ' [EMAIL] ', text)
    
    # Remove phone numbers using unified PHONE_PATTERNS
    for phone_pat in PHONE_PATTERNS:
        text = re.sub(phone_pat, ' [PHONE] ', text)
    
    # Remove URLs
    text = re.sub(r'http\S+|www\.\S+', ' [URL] ', text)
    
    # Mask gender words
    text = re.sub(r'\b(Male|Female|male|female|M|F|Man|Woman|man|woman)\b', ' [GENDER] ', text)
    
    # Remove date of birth patterns
    text = re.sub(r'(?:date\s+of\s+birth|dob|d\.o\.b|born\s+(?:on|in))\s*:?\s*\S+', ' [DOB] ', text, flags=re.IGNORECASE)
    text = re.sub(r'\b\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}\b', ' [DATE] ', text)  # DD/MM/YYYY
    
    # Remove age mentions
    text = re.sub(r'\bage\s*:?\s*\d{1,2}\b', ' [AGE] ', text, flags=re.IGNORECASE)
    text = re.sub(r'\b\d{1,2}\s*years?\s*old\b', ' [AGE] ', text, flags=re.IGNORECASE)
    
    # Remove religion
    religions = r'(?:Hindu|Muslim|Christian|Sikh|Buddhist|Jain|Jewish|Catholic|Protestant|Islam)'
    text = re.sub(r'\b(?:religion|faith)\s*:?\s*' + religions, ' [RELIGION] ', text, flags=re.IGNORECASE)
    text = re.sub(r'\b' + religions + r'\b', ' [RELIGION] ', text, flags=re.IGNORECASE)
    
    # Remove nationality
    text = re.sub(r'\b(?:nationality|citizenship)\s*:?\s*\w+', ' [NATIONALITY] ', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(?:Indian|American|British|Canadian|Australian)\b', ' [NATIONALITY] ', text, flags=re.IGNORECASE)
    
    # Remove marital status
    text = re.sub(r'\b(?:marital\s+status|married|unmarried|single|divorced|widowed)\b', ' [MARITAL] ', text, flags=re.IGNORECASE)
    
    # Remove street addresses
    street_types = r'(?:St|Street|Ave|Avenue|Blvd|Boulevard|Rd|Road|Dr|Drive|Ln|Lane|Way|Ct|Court|Pl|Place|Cir|Circle)'
    text = re.sub(r'\b\d+\s+[A-Za-z]+\s+' + street_types + r'\.?\b', ' [ADDRESS] ', text, flags=re.IGNORECASE)
    
    # Remove zip codes (US: 12345, India: 6-digit)
    text = re.sub(r'\b\d{5,6}(?:-\d{4})?\b', ' [ZIP] ', text)
    
    # Remove city, state patterns
    text = re.sub(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*[A-Z]{2}\b', ' [LOCATION] ', text)
    
    # Remove Aadhaar numbers (Indian ID: 12 digits with spaces)
    text = re.sub(r'\b\d{4}\s+\d{4}\s+\d{4}\b', ' [AADHAAR] ', text)
    
    # Remove PAN numbers (Indian: ABCDE1234F)
    text = re.sub(r'\b[A-Z]{5}\d{4}[A-Z]\b', ' [PAN] ', text)
    
    # Simple header removal (first line if it looks like a name)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if lines:
        first = lines[0]
        if 1 <= len(first.split()) <= 4 and first == first.title():
            lines[0] = "[CANDIDATE NAME]"
        text = "\n".join(lines)
    
    # NLP-based name removal — HEADER SECTION ONLY
    # Uses detect_sections() to find text before first section header.
    # This avoids corrupting names in body text like "Java by James Gosling"
    nlp = _get_nlp()
    if nlp:
        try:
            sections = detect_sections(text)
            header_text = sections.get('header', '')
            if header_text:
                doc = nlp(header_text)
                # Replace by span indices (reverse order to preserve positions)
                replacements = []
                for ent in doc.ents:
                    if ent.label_ == 'PERSON':
                        replacements.append((ent.start_char, ent.end_char))
                # Apply replacements to the header portion of the full text
                for start, end in reversed(replacements):
                    text = text[:start] + '[NAME]' + text[end:]
        except Exception:
            pass
    
    # Compact multiple whitespace
    text = re.sub(r'\s{2,}', ' ', text)
    
    return text


# =============================================================================
# Skill Extraction
# =============================================================================

def extract_skills(text: str) -> List[str]:
    """
    Extract skills from resume text using keyword matching
    
    Args:
        text: Resume text content
    
    Returns:
        list: Extracted skills (flat list for backward compatibility)
    """
    if not isinstance(text, str):
        return []
    
    text_lower = text.lower()
    found_skills = set()
    
    # Check for each skill in the database
    for skill in SKILL_DATABASE:
        # Use word boundaries for exact matches
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            found_skills.add(skill)
    
    # spaCy NER: CONFIRM database skills only, never add raw entities
    # This prevents "Google", "Infosys", "iPhone" from being added as skills
    nlp = _get_nlp()
    if nlp:
        try:
            doc = nlp(text[:10000])  # Limit text length for performance
            for ent in doc.ents:
                if ent.label_ in ['ORG', 'PRODUCT', 'LANGUAGE']:
                    ent_lower = ent.text.lower().strip()
                    # Only add if it matches something in our curated database
                    if ent_lower in SKILL_DATABASE:
                        found_skills.add(ent_lower)
        except Exception as e:
            logger.warning(f"NER extraction failed: {e}")
    
    return sorted(list(found_skills))


def extract_skills_categorized(text: str) -> List[Dict[str, str]]:
    """
    Extract and categorize skills from resume text
    
    Args:
        text: Resume text content
    
    Returns:
        list: List of dicts with 'name' and 'category' keys
    """
    text_lower = text.lower()
    found_skills = []
    seen = set()
    
    for category, skills in SKILL_CATEGORIES.items():
        for skill in skills:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                key = skill.lower()
                if key not in seen:
                    seen.add(key)
                    found_skills.append({
                        'name': skill.title(),
                        'category': category
                    })
    
    return found_skills


# =============================================================================
# Experience Extraction
# =============================================================================

def extract_experience_years(text: str) -> int:
    """
    Extract years of experience from resume text.
    Uses dateparser for smart date-range calculation with regex fallback.
    
    Args:
        text: Resume text
    
    Returns:
        int: Estimated years of experience
    """
    if not isinstance(text, str):
        return 0
    
    text_lower = text.lower()
    max_years = 0
    
    # Strategy 1: Explicit mentions like "5 years of experience"
    explicit_patterns = [
        r'(\d+)\s*(?:\+)?\s*years?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)\s*years?',
        r'(\d+)\s*years?\s+in\s+\w+',
        r'(\d+)\+?\s*yrs?\s+(?:of\s+)?experience'
    ]
    
    for pattern in explicit_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                years = int(match)
                max_years = max(max_years, years)
            except ValueError:
                pass
    
    # Strategy 2: Date-range calculation using dateparser
    # With OVERLAP MERGING — prevents double-counting concurrent jobs
    date_range_pattern = r'([A-Za-z]{3,9}\.?\s*\'?\d{2,4})\s*[-\u2013\u2014to]+\s*(present|current|now|till\s+date|[A-Za-z]{3,9}\.?\s*\'?\d{2,4})'
    date_ranges = re.findall(date_range_pattern, text, re.IGNORECASE)
    
    parsed_intervals = []
    
    for start_str, end_str in date_ranges:
        try:
            start_date = None
            end_date = None
            
            if DATEPARSER_AVAILABLE:
                start_date = dateparser.parse(start_str.strip())
            
            if end_str.strip().lower() in ('present', 'current', 'now', 'till date'):
                end_date = datetime.now()
            elif DATEPARSER_AVAILABLE:
                end_date = dateparser.parse(end_str.strip())
            
            if start_date and end_date and end_date > start_date:
                parsed_intervals.append((start_date, end_date))
        except Exception:
            continue
    
    # Merge overlapping intervals, then sum
    if parsed_intervals:
        merged = _merge_date_intervals(parsed_intervals)
        total_years = sum((end - start).days / 365.25 for start, end in merged)
        computed = int(round(total_years))
        
        # TRUST ALGORITHM OVER TEXT: If we found valid dates, use them.
        # This blocks "Fake Experience Inflation" where candidates claim "10 years" 
        # but only list 2 years of work history.
        return computed

    # Fallback: exact keyword matching (only if no dates found)
    return max_years


def extract_skills_context_aware(text: str) -> List[Dict[str, str]]:
    """
    Context-aware skill extraction: only extracts skills from relevant sections.
    Skills found in HOBBIES, DECLARATION, or REFERENCES are ignored.
    
    Args:
        text: Full resume text
    
    Returns:
        list: List of dicts with 'name', 'category', 'section' keys
    """
    sections = detect_sections(text)
    found_skills = []
    seen = set()
    
    for section_name, section_text in sections.items():
        # Skip sections where skills should be ignored
        if section_name in SKILL_IGNORE_SECTIONS:
            continue
        
        section_lower = section_text.lower()
        
        for category, skills in SKILL_CATEGORIES.items():
            for skill in skills:
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, section_lower):
                    key = skill.lower()
                    if key not in seen:
                        seen.add(key)
                        found_skills.append({
                            'name': skill.title(),
                            'category': category,
                            'section': section_name
                        })
    
    return found_skills


# =============================================================================
# Simple Parse Resume Function (backward compatible)
# =============================================================================

def parse_resume(file_data, filename: str) -> Dict[str, Any]:
    """
    Complete resume parsing pipeline (simple version for backward compatibility)
    
    Args:
        file_data: Binary file data
        filename: Original filename
    
    Returns:
        dict: Parsed resume data with text, skills, experience, anonymized_text
    """
    # Extract text
    text = extract_text_from_file(file_data, filename)
    
    # Extract skills
    skills = extract_skills(text)
    
    # Extract experience years
    experience_years = extract_experience_years(text)
    
    return {
        'raw_text': text,
        'skills': skills,
        'experience_years': experience_years,
        'anonymized_text': anonymize_text(text)
    }


# =============================================================================
# ResumeParser Class - Full-featured parser with all capabilities
# =============================================================================

class ResumeParser:
    """
    Advanced resume parser with NLP capabilities - Entrepreneur Edition v4.0
    
    Supports multiple formats and intelligent information extraction.
    This is the main class for comprehensive resume parsing.
    """
    
    # Class-level references to module constants
    SKILL_CATEGORIES = SKILL_CATEGORIES
    DEGREES = DEGREES
    CERTIFICATIONS = CERTIFICATIONS
    
    def __init__(self):
        """Initialize resume parser with spaCy model"""
        self.nlp = _get_nlp()
        self.spacy_available = SPACY_AVAILABLE
        self.pdf_available = PDF_AVAILABLE
        self.pdfplumber_available = PDFPLUMBER_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        self.dateparser_available = DATEPARSER_AVAILABLE
        
        if not self.spacy_available:
            logger.warning(
                "⚠️ ResumeParser initialized WITHOUT spaCy NLP. "
                "Some features like job title extraction will be limited."
            )
    
    def parse_resume(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse resume from file content - Entrepreneur Edition
        
        Args:
            file_content: Binary file content
            filename: Original filename
        
        Returns:
            Parsed resume data with all extracted information
        """
        try:
            # Extract text based on file type
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                text = extract_text_from_pdf(file_content)
            elif file_ext in ['docx', 'doc']:
                text = extract_text_from_docx(file_content)
            elif file_ext == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            logger.info(f"📄 Extracted {len(text)} characters from resume")
            
            # Detect sections for context-aware extraction
            sections = detect_sections(text)
            
            # Parse information
            parsed_data = {
                'raw_text': text,
                'sections': {k: v[:200] + '...' if len(v) > 200 else v for k, v in sections.items()},
                'contact': self._extract_contact_info(text),
                'skills': self._extract_skills(text),
                'skills_context_aware': extract_skills_context_aware(text),
                'experience': self._extract_experience(text),
                'education': self._extract_education(text),
                'projects': self._extract_projects(sections),
                'languages': self._extract_languages(sections),
                'certifications': self._extract_certifications(text),
                'summary': self._generate_summary(text),
                'anonymized_text': anonymize_text(text),
                'parsed_at': datetime.utcnow().isoformat(),
                'parser_version': '4.2'
            }
            
            # Calculate metadata
            parsed_data['metadata'] = {
                'total_skills': len(parsed_data['skills']),
                'total_skills_context_aware': len(parsed_data['skills_context_aware']),
                'experience_years': parsed_data['experience']['total_years'],
                'education_level': self._calculate_education_level(parsed_data['education']),
                'has_certifications': len(parsed_data['certifications']) > 0,
                'has_projects': len(parsed_data['projects']) > 0,
                'sections_detected': list(sections.keys()),
                'spacy_used': self.spacy_available,
                'pdfplumber_used': self.pdfplumber_available,
                'dateparser_used': self.dateparser_available
            }
            
            # Calculate confidence scores
            parsed_data['confidence'] = self._calculate_confidence(parsed_data)
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"❌ Resume parsing failed: {e}")
            return {
                'error': str(e),
                'parsed_at': datetime.utcnow().isoformat(),
                'parser_version': '4.0'
            }
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information"""
        contact = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None,
            'portfolio': None
        }
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Phone — using unified PHONE_PATTERNS (catches Indian + US + international)
        for phone_pat in PHONE_PATTERNS:
            phones = re.findall(phone_pat, text)
            if phones:
                phone_val = phones[0]
                contact['phone'] = ''.join(phone_val) if isinstance(phone_val, tuple) else phone_val
                break  # Take first match
        
        # LinkedIn
        linkedin_pattern = r'linkedin\.com/in/([A-Za-z0-9_-]+)'
        linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            contact['linkedin'] = f"https://linkedin.com/in/{linkedin_matches[0]}"
        
        # GitHub
        github_pattern = r'github\.com/([A-Za-z0-9_-]+)'
        github_matches = re.findall(github_pattern, text, re.IGNORECASE)
        if github_matches:
            contact['github'] = f"https://github.com/{github_matches[0]}"
        
        # Portfolio (generic URL)
        url_pattern = r'https?://(?:www\.)?([A-Za-z0-9_-]+\.[A-Za-z]{2,})'
        urls = re.findall(url_pattern, text)
        if urls:
            portfolio_urls = [url for url in urls if not any(x in url.lower() for x in ['linkedin', 'github', 'gmail', 'yahoo'])]
            if portfolio_urls:
                contact['portfolio'] = f"https://{portfolio_urls[0]}"
        
        return contact
    
    def _extract_skills(self, text: str) -> List[Dict[str, str]]:
        """Extract and categorize skills"""
        return extract_skills_categorized(text)
    
    def _extract_experience(self, text: str) -> Dict[str, Any]:
        """Extract work experience"""
        experience = {
            'positions': [],
            'total_years': extract_experience_years(text)
        }
        
        # Extract job titles using NLP if available
        if self.nlp:
            try:
                doc = self.nlp(text[:10000])  # Limit for performance
                
                job_indicators = ['developer', 'engineer', 'manager', 'analyst', 'designer', 
                                'architect', 'consultant', 'specialist', 'lead', 'senior',
                                'junior', 'intern', 'director', 'coordinator']
                
                for ent in doc.ents:
                    if ent.label_ == 'ORG' or any(indicator in ent.text.lower() for indicator in job_indicators):
                        experience['positions'].append({
                            'title': ent.text,
                            'context': text[max(0, ent.start_char-50):ent.end_char+50]
                        })
            except Exception as e:
                logger.warning(f"NLP experience extraction failed: {e}")
        
        return experience
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        text_lower = text.lower()
        
        for degree in DEGREES:
            pattern = r'\b' + re.escape(degree) + r'\b'
            if re.search(pattern, text_lower):
                for match in re.finditer(pattern, text_lower):
                    start = max(0, match.start() - 100)
                    end = min(len(text), match.end() + 100)
                    context = text[start:end]
                    
                    education.append({
                        'degree': degree.upper(),
                        'context': context.strip()
                    })
        
        # Remove duplicates
        seen_degrees = set()
        unique_education = []
        for edu in education:
            if edu['degree'] not in seen_degrees:
                seen_degrees.add(edu['degree'])
                unique_education.append(edu)
        
        return unique_education
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certifications = []
        text_lower = text.lower()
        
        for cert_pattern in CERTIFICATIONS:
            pattern = r'\b' + re.escape(cert_pattern)
            matches = re.finditer(pattern, text_lower)
            
            for match in matches:
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                cert_text = text[start:end].strip()
                
                if cert_text not in certifications:
                    certifications.append(cert_text)
        
        return certifications[:10]  # Limit to top 10
    
    def _generate_summary(self, text: str) -> str:
        """Generate a brief summary of the resume"""
        lines = text.split('\n')
        
        summary_keywords = ['summary', 'objective', 'profile', 'about']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in summary_keywords) and len(line_lower) < 50:
                summary_lines = []
                for j in range(i+1, min(i+6, len(lines))):
                    if lines[j].strip() and not any(kw in lines[j].lower() for kw in ['experience', 'education', 'skills']):
                        summary_lines.append(lines[j].strip())
                    else:
                        break
                
                if summary_lines:
                    return ' '.join(summary_lines)
        
        # Fallback: first few non-empty lines
        non_empty_lines = [l.strip() for l in lines[:10] if l.strip() and len(l.strip()) > 20]
        return ' '.join(non_empty_lines[:3]) if non_empty_lines else "No summary available"
    
    def _calculate_education_level(self, education: List[Dict]) -> int:
        """Calculate numeric education level"""
        if not education:
            return 0
        
        level_map = {
            'phd': 4, 'ph.d': 4, 'doctorate': 4,
            'masters': 3, 'master': 3, 'mba': 3, 'ms': 3, 'm.s': 3, 'm.tech': 3, 'mtech': 3, 'me': 3, 'm.e': 3,
            'bachelors': 2, 'bachelor': 2, 'bs': 2, 'b.s': 2, 'ba': 2, 'b.a': 2, 'b.tech': 2, 'btech': 2, 'be': 2, 'b.e': 2,
            'associate': 1, 'diploma': 1
        }
        
        max_level = 0
        for edu in education:
            degree = edu['degree'].lower()
            max_level = max(max_level, level_map.get(degree, 0))
        
        return max_level
    
    def _extract_projects(self, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """Extract projects with tech stacks from the PROJECTS section"""
        projects = []
        project_text = sections.get('projects', '')
        
        if not project_text:
            return projects
        
        lines = [l.strip() for l in project_text.split('\n') if l.strip()]
        current_project = None
        
        for line in lines:
            # Detect project name: lines starting with bullet, dash, or being short title-case
            is_title = (
                line.startswith(('-', '•', '*', '▪')) or
                (len(line.split()) <= 8 and not line.endswith('.'))
            )
            
            if is_title:
                if current_project:
                    projects.append(current_project)
                
                # Clean the line
                name = re.sub(r'^[-•*▪]\s*', '', line).strip()
                
                # Try to extract tech stack from parentheses or after dash
                tech_match = re.search(r'[(\[](.*?)[)\]]', name)
                tech_stack = []
                if tech_match:
                    tech_str = tech_match.group(1)
                    tech_stack = [t.strip() for t in re.split(r'[,;/|]', tech_str) if t.strip()]
                    name = name[:tech_match.start()].strip(' -–—:')
                
                current_project = {
                    'name': name,
                    'tech_stack': tech_stack,
                    'description': ''
                }
            elif current_project:
                # Add description lines
                current_project['description'] += ' ' + line
                
                # Also scan description for tech skills
                if not current_project['tech_stack']:
                    line_lower = line.lower()
                    for skill in SKILL_DATABASE:
                        if re.search(r'\b' + re.escape(skill) + r'\b', line_lower):
                            current_project['tech_stack'].append(skill.title())
        
        if current_project:
            projects.append(current_project)
        
        # Clean up descriptions
        for p in projects:
            p['description'] = p['description'].strip()[:300]
            p['tech_stack'] = list(set(p['tech_stack']))[:10]
        
        return projects[:10]  # Limit to 10 projects
    
    def _extract_languages(self, sections: Dict[str, str]) -> List[Dict[str, str]]:
        """Extract language proficiency from LANGUAGES section"""
        languages = []
        lang_text = sections.get('languages', '')
        
        if not lang_text:
            return languages
        
        # Common proficiency levels
        proficiency_levels = [
            'native', 'fluent', 'proficient', 'advanced', 'intermediate',
            'basic', 'beginner', 'elementary', 'conversational',
            'c2', 'c1', 'b2', 'b1', 'a2', 'a1',
            'mother tongue', 'professional', 'working'
        ]
        
        lines = [l.strip() for l in lang_text.split('\n') if l.strip()]
        
        for line in lines:
            cleaned = re.sub(r'^[-•*▪]\s*', '', line).strip()
            
            # Match patterns like "English - Fluent" or "Hindi (Native)"
            lang_match = re.match(
                r'([A-Za-z]+(?:\s+[A-Za-z]+)?)\s*[-:–(]\s*([^)]+)',
                cleaned
            )
            
            if lang_match:
                lang_name = lang_match.group(1).strip()
                prof = lang_match.group(2).strip().rstrip(')')
                
                # Validate it's a real proficiency level
                prof_lower = prof.lower()
                is_valid = any(level in prof_lower for level in proficiency_levels)
                
                languages.append({
                    'language': lang_name,
                    'proficiency': prof if is_valid else 'Unknown'
                })
            else:
                # Simple language name without proficiency
                words = cleaned.split()
                if words and len(words) <= 3:
                    languages.append({
                        'language': cleaned,
                        'proficiency': 'Unknown'
                    })
        
        return languages[:10]
    
    def _calculate_confidence(self, parsed_data: Dict) -> Dict[str, float]:
        """
        Calculate per-field confidence scores (0.0-1.0).
        Based on extraction quality indicators.
        """
        confidence = {}
        
        # Email: high confidence if valid format found
        contact = parsed_data.get('contact', {})
        confidence['email'] = 0.95 if contact.get('email') else 0.0
        
        # Phone: medium-high if found
        confidence['phone'] = 0.85 if contact.get('phone') else 0.0
        
        # Skills: PRIORITIZE context-aware to block invisible stuffing
        skills = parsed_data.get('skills', [])
        ca_skills = parsed_data.get('skills_context_aware', [])
        
        # Base score on CONTEXT-AWARE skills (harder to fake with white text)
        if len(ca_skills) >= 5:
            confidence['skills'] = 0.9
        elif len(ca_skills) >= 2:
            confidence['skills'] = 0.7
        elif len(skills) > 0:
            # Fallback to raw skills but with lower confidence cap
            confidence['skills'] = 0.4 
        else:
            confidence['skills'] = 0.1
            
        # PENALTY: Suspicious Skill/Experience Ratio (The "Bot" Check)
        # If candidate has 40+ skills but < 3 years experience, likely keyword stuffing
        exp_years = parsed_data.get('experience', {}).get('total_years', 0)
        if len(skills) > 40 and exp_years < 3:
            confidence['skills'] *= 0.1  # Heavy penalty
            confidence['overall_deduction'] = "Suspiciously high skill count for low experience"
        
        # Experience: higher if dateparser found ranges
        if exp_years > 0:
            confidence['experience'] = 0.8
            if DATEPARSER_AVAILABLE:
                confidence['experience'] = 0.9
        else:
            confidence['experience'] = 0.2
        
        # Education: based on degree detection
        education = parsed_data.get('education', [])
        if education:
            confidence['education'] = 0.85
        else:
            confidence['education'] = 0.1
        
        # Projects: based on count
        projects = parsed_data.get('projects', [])
        if projects:
            confidence['projects'] = 0.7
            # Boost if tech stack extracted
            if any(p.get('tech_stack') for p in projects):
                 confidence['projects'] = 0.9
        else:
            confidence['projects'] = 0.0
        
        # Languages
        languages = parsed_data.get('languages', [])
        if languages:
            confidence['languages'] = 0.75
        else:
            confidence['languages'] = 0.0
        
        # Summary
        summary = parsed_data.get('summary', '')
        if summary and summary != "No summary available":
            confidence['summary'] = 0.7
        else:
            confidence['summary'] = 0.2
        
        # Overall: weighted average
        weights = {
            'email': 0.1, 'phone': 0.05, 'skills': 0.35,  # Increased skill weight
            'experience': 0.25, 'education': 0.10,
            'projects': 0.1, 'languages': 0.05
        }
        total = sum(confidence.get(k, 0) * w for k, w in weights.items())
        confidence['overall'] = round(total, 3)
        
        return confidence
    
    def calculate_job_match(self, parsed_resume: Dict, job_requirements: Dict) -> Dict[str, Any]:
        """
        Calculate how well resume matches job requirements
        
        Args:
            parsed_resume: Parsed resume data
            job_requirements: Job requirements dict with skills, experience, education
        
        Returns:
            Match score and breakdown
        """
        match_result = {
            'overall_score': 0,
            'skills_match': 0,
            'experience_match': 0,
            'education_match': 0,
            'matched_skills': [],
            'missing_skills': []
        }
        
        # Skills matching
        if 'required_skills' in job_requirements:
            required_skills = [s.lower() for s in job_requirements['required_skills']]
            candidate_skills = [s['name'].lower() for s in parsed_resume.get('skills', [])]
            
            matched = [s for s in required_skills if s in candidate_skills]
            missing = [s for s in required_skills if s not in candidate_skills]
            
            match_result['matched_skills'] = matched
            match_result['missing_skills'] = missing
            match_result['skills_match'] = (len(matched) / len(required_skills) * 100) if required_skills else 100
        
        # Experience matching
        if 'min_experience_years' in job_requirements:
            required_years = job_requirements['min_experience_years']
            candidate_years = parsed_resume.get('experience', {}).get('total_years', 0)
            
            if candidate_years >= required_years:
                match_result['experience_match'] = 100
            else:
                match_result['experience_match'] = (candidate_years / required_years * 100) if required_years > 0 else 0
        
        # Education matching
        if 'min_education_level' in job_requirements:
            required_level = job_requirements['min_education_level']
            candidate_level = parsed_resume.get('metadata', {}).get('education_level', 0)
            
            if candidate_level >= required_level:
                match_result['education_match'] = 100
            else:
                match_result['education_match'] = (candidate_level / required_level * 100) if required_level > 0 else 0
        
        # Calculate overall score (weighted average)
        weights = {'skills': 0.5, 'experience': 0.3, 'education': 0.2}
        match_result['overall_score'] = (
            match_result['skills_match'] * weights['skills'] +
            match_result['experience_match'] * weights['experience'] +
            match_result['education_match'] * weights['education']
        )
        
        return match_result


# =============================================================================
# Singleton Pattern
# =============================================================================

_resume_parser = None


def get_resume_parser() -> ResumeParser:
    """Get global resume parser instance (singleton)"""
    global _resume_parser
    if _resume_parser is None:
        _resume_parser = ResumeParser()
    return _resume_parser


# =============================================================================
# Exports
# =============================================================================

__all__ = [
    # Main class
    'ResumeParser',
    'get_resume_parser',
    
    # Text extraction functions
    'extract_text_from_file',
    'extract_text_from_pdf', 
    'extract_text_from_docx',
    
    # Skill extraction
    'extract_skills',
    'extract_skills_categorized',
    
    # Other utilities
    'anonymize_text',
    'extract_experience_years',
    'parse_resume',
    
    # Constants
    'SKILL_DATABASE',
    'SKILL_CATEGORIES',
    'DEGREES',
    'CERTIFICATIONS',
    
    # Availability flags
    'PDF_AVAILABLE',
    'DOCX_AVAILABLE',
    'SPACY_AVAILABLE',
]
