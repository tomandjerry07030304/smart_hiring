"""
Resume Parser - Production Version
Extracts text and skills from PDF/DOCX resumes
"""
import re
import io
import logging

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available - PDF parsing disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX parsing disabled")

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    NLP_AVAILABLE = True
except Exception:
    NLP_AVAILABLE = False
    nlp = None
    logging.warning("spaCy not available - using keyword extraction")

def _get_nlp():
    """Get spaCy NLP model if available"""
    return nlp if NLP_AVAILABLE else None

def extract_text_from_pdf(file_data):
    """
    Extract text from PDF file
    
    Args:
        file_data: Binary PDF file data or file-like object
    
    Returns:
        str: Extracted text content
    """
    if not PDF_AVAILABLE:
        return "PDF parsing not available. Please install PyPDF2."
    
    try:
        # Handle both bytes and file-like objects
        if isinstance(file_data, bytes):
            pdf_file = io.BytesIO(file_data)
        else:
            pdf_file = file_data
        
        reader = PdfReader(pdf_file)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = "\n".join(text_parts)
        return full_text if full_text.strip() else "No text could be extracted from PDF"
    
    except Exception as e:
        logging.error(f"PDF extraction error: {e}")
        return f"Error extracting PDF: {str(e)}"

def extract_text_from_docx(file_data):
    """
    Extract text from DOCX file
    
    Args:
        file_data: Binary DOCX file data or file-like object
    
    Returns:
        str: Extracted text content
    """
    if not DOCX_AVAILABLE:
        return "DOCX parsing not available. Please install python-docx."
    
    try:
        # Handle both bytes and file-like objects
        if isinstance(file_data, bytes):
            docx_file = io.BytesIO(file_data)
        else:
            docx_file = file_data
        
        doc = Document(docx_file)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        return full_text if full_text.strip() else "No text could be extracted from DOCX"
    
    except Exception as e:
        logging.error(f"DOCX extraction error: {e}")
        return f"Error extracting DOCX: {str(e)}"

def extract_text_from_file(file_data, filename):
    """Extract text from uploaded file - stub version"""
    name = filename.lower()
    
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_data)
    elif name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx(file_data)
    else:
        # Treat as plain text
        try:
            return file_data.decode('utf-8', errors='ignore')
        except Exception:
            return str(file_data)

def anonymize_text(text):
    """Remove PII from text using basic regex - simplified version"""
    if not isinstance(text, str):
        return ""
    
    # Remove emails
    text = re.sub(r'\S+@\S+', ' [EMAIL] ', text)
    
    # Remove phone numbers (various formats)
    text = re.sub(r'\+?\d[\d\-\s()]{6,}\d', ' [PHONE] ', text)
    
    # Remove URLs
    text = re.sub(r'http\S+|www\.\S+', ' [URL] ', text)
    
    # Mask gender words
    text = re.sub(r'\b(Male|Female|male|female|M|F|Man|Woman|man|woman)\b', ' [GENDER] ', text)
    
    # Simple header removal (first line if it looks like a name)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if lines:
        first = lines[0]
        if 1 <= len(first.split()) <= 4 and first == first.title():
            lines[0] = "[REDACTED HEADER]"
        text = "\n".join(lines)
    
    # Compact whitespace
    text = re.sub(r'\s{2,}', ' ', text)
    return text


# Comprehensive skill database (200+ skills)
SKILL_DATABASE = {
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift',
    'kotlin', 'go', 'rust', 'scala', 'r', 'matlab', 'perl', 'shell', 'bash',
    
    # Web Technologies
    'html', 'css', 'react', 'angular', 'vue.js', 'node.js', 'express', 'django', 'flask',
    'spring', 'asp.net', 'jquery', 'bootstrap', 'tailwind', 'sass', 'webpack', 'vite',
    
    # Databases
    'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'cassandra',
    'dynamodb', 'elasticsearch', 'neo4j', 'firebase',
    
    # Cloud & DevOps
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'gitlab', 'github actions',
    'terraform', 'ansible', 'ci/cd', 'devops', 'linux', 'unix',
    
    # Data Science & ML
    'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
    'pandas', 'numpy', 'nlp', 'computer vision', 'data analysis', 'statistics',
    'tableau', 'power bi', 'spark', 'hadoop', 'kafka',
    
    # Mobile Development
    'android', 'ios', 'react native', 'flutter', 'xamarin', 'swift ui',
    
    # Tools & Frameworks
    'git', 'jira', 'agile', 'scrum', 'restful api', 'graphql', 'microservices',
    'api development', 'testing', 'junit', 'pytest', 'selenium', 'jest',
    
    # Soft Skills
    'communication', 'leadership', 'team player', 'problem solving', 'critical thinking',
    'project management', 'time management', 'collaboration',
    
    # Security
    'cybersecurity', 'penetration testing', 'oauth', 'jwt', 'encryption', 'ssl/tls',
    
    # Design
    'ui/ux', 'figma', 'adobe xd', 'photoshop', 'illustrator', 'graphic design',
    
    # Business
    'product management', 'business analysis', 'data analysis', 'excel', 'powerpoint'
}


def extract_skills(text):
    """
    Extract skills from resume text using keyword matching
    
    Args:
        text: Resume text content
    
    Returns:
        list: Extracted skills
    """
    if not isinstance(text, str):
        return []
    
    text_lower = text.lower()
    found_skills = set()
    
    # Check for each skill in the database
    for skill in SKILL_DATABASE:
        # Use word boundaries for exact matches
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            found_skills.add(skill)
    
    # If spaCy available, use NER for additional extraction
    if NLP_AVAILABLE and nlp:
        try:
            doc = nlp(text[:10000])  # Limit text length for performance
            for ent in doc.ents:
                if ent.label_ in ['ORG', 'PRODUCT', 'LANGUAGE']:
                    skill_text = ent.text.lower().strip()
                    if len(skill_text) > 2:  # Avoid very short matches
                        found_skills.add(skill_text)
        except Exception as e:
            logging.warning(f"NER extraction failed: {e}")
    
    return sorted(list(found_skills))


def parse_resume(file_data, filename):
    """
    Complete resume parsing pipeline
    
    Args:
        file_data: Binary file data
        filename: Original filename
    
    Returns:
        dict: Parsed resume data with text, skills, etc.
    """
    # Extract text
    text = extract_text_from_file(file_data, filename)
    
    # Extract skills
    skills = extract_skills(text)
    
    # Extract basic info (years of experience)
    experience_years = extract_experience_years(text)
    
    return {
        'raw_text': text,
        'skills': skills,
        'experience_years': experience_years,
        'anonymized_text': anonymize_text(text)
    }


def extract_experience_years(text):
    """
    Extract years of experience from resume
    
    Args:
        text: Resume text
    
    Returns:
        int: Estimated years of experience
    """
    if not isinstance(text, str):
        return 0
    
    text_lower = text.lower()
    
    # Look for explicit mentions like "5 years of experience"
    patterns = [
        r'(\d+)\s*(?:\+)?\s*years?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)\s*years?',
        r'(\d+)\s*years?\s+in\s+\w+',
    ]
    
    max_years = 0
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                years = int(match)
                max_years = max(max_years, years)
            except ValueError:
                pass
    
    return max_years
